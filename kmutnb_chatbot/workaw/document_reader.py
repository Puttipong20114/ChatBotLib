# document_reader.py
from typing import List, Optional
import os

# -------- Optional deps (ไม่มีก็ข้ามได้) --------
# DOCX
try:
    import docx  # python-docx
except Exception:
    docx = None

# PDF ตัวเลือกที่ 1: pdfplumber (ดีในหลายเคส)
try:
    import pdfplumber
except Exception:
    pdfplumber = None

# PDF ตัวเลือกที่ 2: pypdf (fallback)
try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

# PDF ตัวเลือกที่ 3: pdfminer.six (fallback เพิ่ม)
try:
    from io import StringIO
    from pdfminer.high_level import extract_text_to_fp
except Exception:
    extract_text_to_fp = None  # ไม่มี pdfminer ก็ข้าม

# OCR/Visual อ่านไฟล์ด้วย Gemini (fallback ขั้นสุด)
try:
    import google.generativeai as genai
except Exception:
    genai = None  # ถ้าไม่มีก็ไม่ใช้ OCR

MAX_CHARS = 120_000  # กันข้อความยาวเกินส่งเข้าโมเดลหลัก
MIN_OK_LEN = 80      # ยาวเกินนี้ถือว่าอ่านได้จริง

# ---------------- Utils ----------------
def _clean_text(text: str) -> str:
    if not text:
        return ""
    t = text.replace("\x00", "").replace("\r", "")
    # strip ทีละบรรทัด + ตัดบรรทัดว่าง
    return "\n".join(line.strip() for line in t.split("\n") if line.strip())

# ---------------- Readers ----------------
def _read_docx(path: str) -> str:
    if docx is None:
        return "Error: python-docx is not installed. Please install it with: pip install python-docx"
    try:
        d = docx.Document(path)
        parts: List[str] = []
        for p in d.paragraphs:
            parts.append(p.text or "")
        # ตาราง (ถ้ามี)
        for table in d.tables:
            for row in table.rows:
                parts.append(" | ".join((cell.text or "") for cell in row.cells))
        text = "\n".join(parts)
        return _clean_text(text)[:MAX_CHARS]
    except Exception as e:
        return f"Error: cannot read DOCX -> {e}"

def _read_pdf_plumber(path: str) -> Optional[str]:
    if pdfplumber is None:
        return None
    try:
        parts: List[str] = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                txt = page.extract_text() or ""
                parts.append(txt)
        text = _clean_text("\n".join(parts))
        return text[:MAX_CHARS]
    except Exception:
        return None  # ให้ไปลองวิธีอื่น

def _read_pdf_pypdf(path: str) -> Optional[str]:
    if PdfReader is None:
        return None
    try:
        reader = PdfReader(path)
        parts: List[str] = []
        for page in reader.pages:
            txt = page.extract_text() or ""
            parts.append(txt)
        text = _clean_text("\n".join(parts))
        return text[:MAX_CHARS]
    except Exception:
        return None

def _read_pdf_pdfminer(path: str) -> Optional[str]:
    if extract_text_to_fp is None:
        return None
    try:
        output = StringIO()
        with open(path, "rb") as f:
            extract_text_to_fp(f, output, laparams=None)
        text = _clean_text(output.getvalue())
        return text[:MAX_CHARS]
    except Exception:
        return None

def _summarize_with_gemini(file_path: str) -> Optional[str]:
    """
    Fallback ขั้นสุด: อัปโหลด PDF ให้ Gemini 1.5 อ่าน/มองตรง ๆ
    ใช้ได้แม้เป็นสแกนรูป (OCR/vision) แล้วสรุปเป็น knowledge สำหรับ chatbot
    NOTE: genai.configure(...) ทำแล้วใน app.py ห้ามตั้งซ้ำที่นี่
    """
    if genai is None:
        return None
    try:
        uploaded = genai.upload_file(file_path)  # อัปโหลดไฟล์
        model = genai.GenerativeModel("gemini-1.5-flash")
        prompt = (
            "คุณคือผู้ช่วยสรุปเอกสารหอสมุด KMUTNB. "
            "สกัดเฉพาะข้อมูลข้อเท็จจริงที่จำเป็นต่อการตอบคำถามผู้ใช้ เช่น "
            "เวลาทำการ ขั้นตอนยืม-คืน ค่าปรับ วิธีติดต่อ บริการ และ FAQ. "
            "ตอบเป็นหัวข้อ bullet สั้น กระชับ อ้างอิงจากไฟล์เท่านั้น."
        )
        resp = model.generate_content([prompt, uploaded])
        text = (getattr(resp, "text", "") or "").strip()
        if not text:
            return None
        text = _clean_text(text)
        return text[:MAX_CHARS]
    except Exception:
        return None

def _read_pdf(path: str) -> str:
    # ลองตามลำดับ: pdfplumber → pypdf → pdfminer → Gemini OCR
    for reader in (_read_pdf_plumber, _read_pdf_pypdf, _read_pdf_pdfminer):
        try:
            text = reader(path)
            if text and len(text.strip()) >= MIN_OK_LEN:
                return text
        except Exception:
            pass

    # สแกน/ไม่มี text layer → ลอง OCR/vision ผ่าน Gemini
    ocr_summary = _summarize_with_gemini(path)
    if ocr_summary and len(ocr_summary) >= 40:
        return "## OCR/AI Summary\n" + ocr_summary

    # ไปต่อไม่ไหวจริง ๆ
    return "Error: cannot extract text from PDF (tried pdfplumber, pypdf, pdfminer, and Gemini OCR fallback)."

# ---------------- Public API ----------------
def get_kmutnb_summary(file_path: str) -> str:
    """
    อ่านไฟล์ dataset (.pdf หรือ .docx) แล้วคืนค่าเป็นข้อความแบบสรุป (plain text)
    ถ้าเกิดปัญหา คืนค่า string ที่ขึ้นต้นด้วย 'Error:' เพื่อให้โค้ดหลักจัดการได้
    """
    if not os.path.exists(file_path):
        return f"Error: file not found -> {file_path}"

    ext = os.path.splitext(file_path)[1].lower()

    try:
        if ext == ".pdf":
            return _read_pdf(file_path)
        elif ext == ".docx":
            return _read_docx(file_path)
        else:
            return "Error: unsupported file type. Please use .pdf or .docx"
    except Exception as e:
        return f"Error: {e}"
